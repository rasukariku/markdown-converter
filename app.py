import os
import re
import io
import tempfile
import platform
import pypandoc
from flask import Flask, request, send_file, render_template
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# Initialize Pandoc
try:
    pypandoc.get_pandoc_version()
except OSError:
    pypandoc.download_pandoc()

app = Flask(__name__)

def _get_buffer_and_cleanup(filepath):
    with open(filepath, 'rb') as f:
        buffer = io.BytesIO(f.read())
    buffer.seek(0)
    try:
        if os.path.exists(filepath): os.unlink(filepath)
    except OSError: pass
    return buffer

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    markdown_content = request.form.get('markdown_content', '')
    file_format = request.form.get('file_format', 'docx')

    if not markdown_content:
        return "Input text cannot be empty.", 400

    # =========================================================================
    # 1. PRE-PROCESSING: Normalize Horizontal Rules using Placeholder
    # =========================================================================
    # Replace all HR markdown syntax with a unique placeholder text
    markdown_content = re.sub(r'<hr\s*/?>', '\n\n[[HR_PLACEHOLDER]]\n\n', markdown_content, flags=re.IGNORECASE)
    markdown_content = re.sub(r'(?m)^\s*(\*{3,}|-{3,}|_{3,})\s*$', '\n\n[[HR_PLACEHOLDER]]\n\n', markdown_content)

    source_text = markdown_content
    source_text_html = source_text.replace('[[HR_PLACEHOLDER]]', '\n\n---\n\n')
    input_format = 'markdown+raw_html'

    # =========================================================================
    # 2. HTML EXPORT
    # =========================================================================
    if file_format == 'html':
        temp_html = tempfile.NamedTemporaryFile(delete=False, suffix='.html')
        temp_html.close()
        try:
            extra_args = [
                '--standalone', 
                '--mathjax=https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js'
            ]
            pypandoc.convert_text(source_text_html, 'html', format=input_format, outputfile=temp_html.name, extra_args=extra_args)
            buffer = _get_buffer_and_cleanup(temp_html.name)
            return send_file(buffer, as_attachment=True, download_name='Markdown_Export.html', mimetype='text/html')
        except Exception as e:
            if os.path.exists(temp_html.name): os.unlink(temp_html.name)
            return f"HTML conversion error: {str(e)}", 500

    # =========================================================================
    # 3. BASE DOCX CONVERSION (via Pandoc)
    # =========================================================================
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_docx.close()

    try:
        pypandoc.convert_text(source_text, 'docx', format=input_format, outputfile=temp_docx.name, extra_args=['--highlight-style=tango'])
        doc = Document(temp_docx.name)

        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        normal_style.font.color.rgb = RGBColor(0, 0, 0)

        try:
            hlink_style = doc.styles['Hyperlink']
            hlink_style.font.name = 'Times New Roman'
            hlink_style.font.color.rgb = RGBColor(0, 0, 0)
            if hlink_style.element.rPr is not None:
                color_el = hlink_style.element.rPr.find(qn('w:color'))
                if color_el is not None and qn('w:themeColor') in color_el.attrib: del color_el.attrib[qn('w:themeColor')]
                rFonts = hlink_style.element.rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme']:
                        if qn(attr) in rFonts.attrib: del rFonts.attrib[qn(attr)]
        except KeyError: pass

        for section in doc.sections:
            section.page_width, section.page_height = Cm(21.0), Cm(29.7)
            section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = Cm(2.54), Cm(2.54), Cm(2.54), Cm(2.54)

        settings = doc.settings.element
        math_pr = settings.find(qn('m:mathPr'))
        if math_pr is None: math_pr = OxmlElement('m:mathPr'); settings.append(math_pr)
        def_jc = math_pr.find(qn('m:defJc'))
        if def_jc is None: def_jc = OxmlElement('m:defJc'); math_pr.append(def_jc)
        def_jc.set(qn('m:val'), 'left')

        compat = settings.find(qn('w:compat'))
        if compat is None: compat = OxmlElement('w:compat'); settings.append(compat)
        compat_setting = OxmlElement('w:compatSetting')
        compat_setting.set(qn('w:name'), 'compatibilityMode')
        compat_setting.set(qn('w:uri'), 'http://schemas.microsoft.com/office/word')
        compat_setting.set(qn('w:val'), '15')
        compat.append(compat_setting)

        # =========================================================================
        # 4. PARAGRAPH FORMATTING LOOP
        # =========================================================================
        paragraphs = list(doc.paragraphs)
        removal_queue = []

        for i, para in enumerate(paragraphs):
            text_clean = para.text.strip()
            style_name = para.style.name

            has_math = bool(para._element.findall('.//' + qn('m:oMath'))) or bool(para._element.findall('.//' + qn('m:oMathPara')))
            has_drawing = bool(para._element.findall('.//' + qn('w:drawing')))
            is_heading = style_name.startswith('Heading')
            is_list = ('List' in style_name or 'Bullet' in style_name or 'Compact' in style_name or bool(para._element.findall('.//' + qn('w:numPr'))))
            is_code = 'Source Code' in style_name or 'Code' in style_name
            is_quote = 'Quote' in style_name or 'Block Text' in style_name
            has_soft_return = '\n' in para.text

            # -----------------------------------------------------------------
            # HORIZONTAL RULE PROCESSING VIA PLACEHOLDER
            # -----------------------------------------------------------------
            is_hr = '[[HR_PLACEHOLDER]]' in text_clean
            
            if is_hr:
                p_el = para._element
                for run in list(para.runs): 
                    p_el.remove(run._element)
                
                pPr = p_el.get_or_add_pPr()
                
                old_bdr = pPr.find(qn('w:pBdr'))
                if old_bdr is not None: pPr.remove(old_bdr)

                pBdr = OxmlElement('w:pBdr')
                top_border = OxmlElement('w:top')
                top_border.set(qn('w:val'), 'single')
                top_border.set(qn('w:sz'), '12') # Ketebalan 1.5pt
                top_border.set(qn('w:space'), '0') # Nol jarak
                top_border.set(qn('w:color'), '000000')
                pBdr.append(top_border)
                pPr.append(pBdr)

                try: para.style = doc.styles['Normal']
                except KeyError: pass

                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                para.paragraph_format.line_spacing = Pt(1)
                
                para.paragraph_format.space_before = Pt(0)     
                para.paragraph_format.space_after = Pt(0)      
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)

                has_border = True
                text_clean = ""
            else:
                has_border = False

            if not text_clean and not has_math and not has_drawing and not is_code and not has_border:
                removal_queue.append(para)
                continue

            if is_code:
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.left_indent = Pt(18)
                para.paragraph_format.first_line_indent = Pt(0)
            else:
                if not is_heading and not is_list and not is_quote and not has_border:
                    para.style = doc.styles['Normal']

                para.paragraph_format.line_spacing = 1.5
                para.paragraph_format.space_before = Pt(0)

                if is_list:
                    next_is_list = False
                    if i + 1 < len(paragraphs):
                        next_para = paragraphs[i + 1]
                        next_is_list = ('List' in next_para.style.name or 'Bullet' in next_para.style.name or 'Compact' in next_para.style.name or bool(next_para._element.findall('.//' + qn('w:numPr'))))
                    para.paragraph_format.space_after = Pt(5) if next_is_list else Pt(8)
                    ilvl_nodes = para._element.findall('.//' + qn('w:ilvl'))
                    level = int(ilvl_nodes[0].get(qn('w:val'))) if ilvl_nodes else 0
                    para.paragraph_format.left_indent = Pt(36 + (level * 36))
                    para.paragraph_format.first_line_indent = Pt(-18)
                elif is_quote:
                    para.paragraph_format.left_indent = Cm(1.5)
                    para.paragraph_format.right_indent = Cm(1.5)
                    para.paragraph_format.space_after = Pt(8)
                else:
                    if not has_border:
                        para.paragraph_format.space_after = Pt(8)
                        para.paragraph_format.left_indent = Pt(0)
                        para.paragraph_format.first_line_indent = Pt(0)

            for run in para.runs:
                if run._element.findall('.//' + qn('m:oMath')): continue

                if is_code:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10.5)
                else:
                    run.font.name = 'Times New Roman'
                    if not is_heading and run.font.size is None: run.font.size = Pt(12)

                    run.font.color.rgb = RGBColor(0, 0, 0)
                    rPr = run._element.get_or_add_rPr()
                    color_el = rPr.find(qn('w:color'))
                    if color_el is not None and qn('w:themeColor') in color_el.attrib: del color_el.attrib[qn('w:themeColor')]

                    if 'Hyperlink' in run.style.name or 'Hyperlink' in style_name: run.font.underline = True

                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None: rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    rFonts.set(qn('w:cs'), 'Times New Roman')
                    for attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme']:
                        if qn(attr) in rFonts.attrib: del rFonts.attrib[qn(attr)]

                rPr = run._element.get_or_add_rPr()
                lang_el = rPr.find(qn('w:lang'))
                if lang_el is None: lang_el = OxmlElement('w:lang'); rPr.append(lang_el)
                lang_el.set(qn('w:val'), 'id-ID')

            if is_heading:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_after = Pt(12)
            elif is_list or is_code or has_soft_return or has_border:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if len(text_clean) > 80 else WD_ALIGN_PARAGRAPH.LEFT

        for para in reversed(removal_queue):
            p = para._element
            if p.getparent() is not None: p.getparent().remove(p)

        # =========================================================================
        # 5. TABLE PROCESSING
        # =========================================================================
        for table in doc.tables:
            try: table.style = 'Table Grid'
            except KeyError:
                try: table.style = 'TableGrid'
                except KeyError: pass

            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        orig_align = para.alignment
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
                        para.paragraph_format.line_spacing = 1.5
                        para.alignment = orig_align if orig_align is not None else WD_ALIGN_PARAGRAPH.LEFT

                        for run in para.runs:
                            if run._element.findall('.//' + qn('m:oMath')): continue
                            run.font.name = 'Times New Roman'
                            if run.font.size is None: run.font.size = Pt(12)
                            if run.font.color.rgb is None and 'Hyperlink' not in run.style.name: run.font.color.rgb = RGBColor(0, 0, 0)

            tbl_element = table._element
            next_element = tbl_element.getnext()
            if next_element is not None and next_element.tag == qn('w:p'):
                next_para = Paragraph(next_element, doc._body)
                next_para.paragraph_format.space_before = Pt(8)

        doc.save(temp_docx.name)

        # =========================================================================
        # 6. PDF GENERATION
        # =========================================================================
        if file_format == 'pdf':
            current_os = platform.system()
            if current_os == 'Windows':
                import win32com.client, pythoncom
                pythoncom.CoInitialize()
                temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf'); temp_pdf.close()
                word = None
                try:
                    word = win32com.client.DispatchEx("Word.Application")
                    word.Visible, word.DisplayAlerts = False, False
                    doc_com = word.Documents.Open(os.path.abspath(temp_docx.name), ReadOnly=True, Visible=False)
                    doc_com.SaveAs(os.path.abspath(temp_pdf.name), FileFormat=17)
                    doc_com.Close(SaveChanges=False)
                    if os.path.exists(temp_docx.name): os.unlink(temp_docx.name)
                    return send_file(_get_buffer_and_cleanup(temp_pdf.name), as_attachment=True, download_name='Markdown_Export.pdf', mimetype='application/pdf')
                except Exception as e:
                    if os.path.exists(temp_pdf.name): os.unlink(temp_pdf.name)
                    if os.path.exists(temp_docx.name): os.unlink(temp_docx.name)
                    return f"Windows PDF Engine Error: {str(e)}", 500
                finally:
                    if word:
                        try: word.Quit()
                        except: pass
                    pythoncom.CoUninitialize()
            else:
                try:
                    from playwright.sync_api import sync_playwright
                    
                    temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
                    temp_pdf.close()
                    
                    extra_args = [
                        '--standalone', 
                        '--mathjax=https://cdn.jsdelivr.net/npm/mathjax@3/es5/tex-mml-chtml.js'
                    ]
                    html_string = pypandoc.convert_text(source_text_html, 'html', format=input_format, extra_args=extra_args)
                    
                    css_injection = '''
                    <style>
                    body { font-family: "Times New Roman", serif; font-size: 16px; line-height: 1.5; text-align: justify; color: black; }
                    h1, h2, h3, h4 { line-height: 1.2; margin-bottom: 8px; text-align: left; font-family: "Times New Roman", serif; }
                    p { margin-bottom: 10px; margin-top: 0; }
                    table { width: 100%; border-collapse: collapse; margin: 15px 0; page-break-inside: avoid; }
                    th, td { border: 1px solid black; padding: 8px; text-align: left; vertical-align: top; }
                    th { font-weight: bold; background-color: #f3f4f6; }
                    blockquote { margin: 10px 20px; padding-left: 10px; border-left: 3px solid #000; font-style: italic; }
                    hr { border: 0; border-top: 1px solid #000; margin: 16px 0; }
                    pre, code { font-family: "Courier New", monospace; font-size: 14px; page-break-inside: avoid; }
                    pre { background: #f4f4f4; padding: 10px; border: 1px solid #ccc; white-space: pre-wrap; }
                    
                    mjx-container { page-break-inside: avoid !important; margin: 6px 0 !important; }
                    </style>
                    '''
                    html_string = html_string.replace('</head>', f'{css_injection}</head>')
                    
                    with sync_playwright() as p:
                        browser = p.chromium.launch(headless=True)
                        page = browser.new_page()
                        
                        page.set_content(html_string, wait_until='networkidle')
                        
                        page.pdf(
                            path=temp_pdf.name,
                            format='A4',
                            margin={'top': '2.54cm', 'right': '2.54cm', 'bottom': '2.54cm', 'left': '2.54cm'},
                            print_background=True
                        )
                        browser.close()
                    
                    if os.path.exists(temp_docx.name): os.unlink(temp_docx.name)
                    return send_file(_get_buffer_and_cleanup(temp_pdf.name), as_attachment=True, download_name='Markdown_Export.pdf', mimetype='application/pdf')
                
                except Exception as e:
                    if os.path.exists(temp_pdf.name): os.unlink(temp_pdf.name)
                    if os.path.exists(temp_docx.name): os.unlink(temp_docx.name)
                    return f"Linux PDF Engine Error: {str(e)}", 500

        # =========================================================================
        # 7. DOCX OUTPUT
        # =========================================================================
        buffer = _get_buffer_and_cleanup(temp_docx.name)
        return send_file(buffer, as_attachment=True, download_name='Markdown_Export.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except Exception as e:
        if os.path.exists(temp_docx.name): os.unlink(temp_docx.name)
        return f"DOCX conversion error: {str(e)}", 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=int(os.environ.get('PORT', 7860)))